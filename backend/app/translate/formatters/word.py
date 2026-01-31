"""
Word 文档格式处理器
支持 .docx 文件的翻译，保持格式
参考开源项目 DocTranslator 的实现
"""
import uuid
import asyncio
from pathlib import Path
from typing import Callable, Optional
from docx import Document
from docx.oxml.ns import qn

from . import BaseFormatter
from ...config import settings


class WordFormatter(BaseFormatter):
    """
    Word 文档处理器

    使用 python-docx 库处理 Word 文档
    保持段落格式、字体样式、表格等
    """

    def translate(
        self,
        source_path: str,
        target_lang: str,
        ai_translator,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> str:
        """
        翻译 Word 文档（同步包装器，调用异步方法）

        Args:
            source_path: 源文件路径
            target_lang: 目标语言
            ai_translator: AI 翻译器实例
            progress_callback: 进度回调函数

        Returns:
            str: 翻译结果文件路径
        """
        # 检查翻译器是否支持并发方法
        if hasattr(ai_translator, 'translate_batch_async_concurrent'):
            # 获取线程数配置
            thread_count = getattr(ai_translator, 'thread_count', 5)

            # 创建新的事件循环并运行异步方法（不关闭循环，避免 AsyncOpenAI 客户端引用错误）
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                return loop.run_until_complete(
                    self.translate_async(source_path, target_lang, ai_translator, thread_count, progress_callback)
                )
            finally:
                # 不关闭循环，让 asyncio 自动管理
                asyncio.set_event_loop(None)
        else:
            # 使用原有同步逻辑
            return self._translate_sync(source_path, target_lang, ai_translator, progress_callback)

    def _translate_sync(
        self,
        source_path: str,
        target_lang: str,
        ai_translator,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> str:
        """
        同步翻译 Word 文档（支持对照模式）

        Args:
            source_path: 源文件路径
            target_lang: 目标语言
            ai_translator: AI 翻译器实例
            progress_callback: 进度回调函数

        Returns:
            str: 翻译结果文件路径
        """
        # 获取显示模式配置（数字映射）
        # 1=替换模式, 2=对照模式, 3=表格对照, 4=双语对照...
        display_mode = getattr(ai_translator, 'display_mode', 1)

        # 加载文档
        doc = Document(source_path)

        # 第一步：读取所有需要翻译的文本（段落级别，不是 run 级别）
        paragraphs_data = []

        # 读取段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                paragraphs_data.append({
                    'type': 'paragraph',
                    'index': para_idx,
                    'text': text,
                    'translated': None,
                    'element': paragraph
                })

        # 读取表格中的文本
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()
                        if text:
                            paragraphs_data.append({
                                'type': 'table',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'para_idx': para_idx,
                                'text': text,
                                'translated': None,
                                'element': paragraph
                            })

        # 总文本数
        total_count = len(paragraphs_data)
        if total_count == 0:
            return self._generate_result_path(source_path, target_lang)

        print(f"📊 文档包含 {total_count} 个段落，开始翻译...")
        print(f"🎨 显示模式: {display_mode}")

        # 第二步：批量翻译
        batch_size = 10

        for i in range(0, total_count, batch_size):
            batch = paragraphs_data[i:i + batch_size]
            batch_texts = [item['text'] for item in batch]

            # 调用 AI 翻译
            translated_batch = ai_translator.translate_batch(
                texts=batch_texts,
                target_lang=target_lang
            )

            # 将翻译结果写回
            for j, item in enumerate(batch):
                if j < len(translated_batch):
                    item['translated'] = translated_batch[j]
                else:
                    item['translated'] = item['text']

            # 更新进度
            if progress_callback:
                current_completed = min(i + batch_size, total_count)
                progress_callback(current_completed, total_count)
                print(f"📊 批次 {i // batch_size + 1} 完成，进度: {int(current_completed / total_count * 100)}%")

        # 第三步：根据显示模式处理文档
        if display_mode == 2:
            # 2=对照模式：原文保留，添加译文在下
            self._write_parallel_mode(paragraphs_data, target_lang)
        elif display_mode == 3:
            # 3=表格对照模式（预留）
            self._write_parallel_mode(paragraphs_data, target_lang)
        elif display_mode == 4:
            # 4=双语对照模式（预留）
            self._write_replace_mode(paragraphs_data)
        else:
            # 1=替换模式：直接替换原文
            self._write_replace_mode(paragraphs_data)

        # 设置中文字体
        self._set_chinese_font(doc, target_lang)

        # 保存结果
        result_path = self._generate_result_path(source_path, target_lang)
        doc.save(result_path)

        return result_path

    async def translate_async(
        self,
        source_path: str,
        target_lang: str,
        ai_translator,
        thread_count: int = 5,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> str:
        """
        异步翻译 Word 文档（支持并发和对照模式）

        Args:
            source_path: 源文件路径
            target_lang: 目标语言
            ai_translator: AI 翻译器实例
            thread_count: 并发线程数
            progress_callback: 进度回调函数

        Returns:
            str: 翻译结果文件路径
        """
        # 获取显示模式配置（数字映射）
        # 1=替换模式, 2=对照模式, 3=表格对照, 4=双语对照...
        display_mode = getattr(ai_translator, 'display_mode', 1)

        # 加载文档
        doc = Document(source_path)

        # 第一步：读取所有需要翻译的文本（段落级别）
        paragraphs_data = []

        # 读取段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                paragraphs_data.append({
                    'type': 'paragraph',
                    'index': para_idx,
                    'text': text,
                    'translated': None,
                    'element': paragraph
                })

        # 读取表格中的文本
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()
                        if text:
                            paragraphs_data.append({
                                'type': 'table',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'para_idx': para_idx,
                                'text': text,
                                'translated': None,
                                'element': paragraph
                            })

        # 总文本数
        total_count = len(paragraphs_data)
        if total_count == 0:
            return self._generate_result_path(source_path, target_lang)

        print(f"📊 文档包含 {total_count} 个段落，开始异步翻译（{thread_count} 线程）...")
        print(f"🎨 显示模式: {display_mode}")

        # 第二步：异步批量翻译（并发模式）
        batch_size = 20

        # 检查翻译器是否支持并发方法
        if hasattr(ai_translator, 'translate_batch_async_concurrent'):
            # 使用并发翻译
            for i in range(0, total_count, batch_size):
                batch = paragraphs_data[i:i + batch_size]
                batch_texts = [item['text'] for item in batch]

                def batch_progress_callback(batch_current: int, batch_total: int):
                    if progress_callback:
                        global_current = i + batch_current
                        progress_callback(global_current, total_count)

                translated_batch = await ai_translator.translate_batch_async_concurrent(
                    texts=batch_texts,
                    target_lang=target_lang,
                    max_concurrency=thread_count,
                    progress_callback=batch_progress_callback
                )

                for j, item in enumerate(batch):
                    if j < len(translated_batch):
                        item['translated'] = translated_batch[j]
                    else:
                        item['translated'] = item['text']
        else:
            # 降级到普通异步翻译
            for i in range(0, total_count, batch_size):
                batch = paragraphs_data[i:i + batch_size]
                batch_texts = [item['text'] for item in batch]

                translated_batch = await ai_translator.translate_batch_async(
                    texts=batch_texts,
                    target_lang=target_lang
                )

                for j, item in enumerate(batch):
                    if j < len(translated_batch):
                        item['translated'] = translated_batch[j]
                    else:
                        item['translated'] = item['text']

                if progress_callback:
                    progress_callback(min(i + batch_size, total_count), total_count)

        # 第三步：根据显示模式处理文档
        if display_mode == 2:
            # 2=对照模式
            self._write_parallel_mode(paragraphs_data, target_lang)
        elif display_mode == 3:
            # 3=表格对照模式（预留）
            self._write_parallel_mode(paragraphs_data, target_lang)
        elif display_mode == 4:
            # 4=双语对照模式（预留）
            self._write_replace_mode(paragraphs_data)
        else:
            # 1=替换模式
            self._write_replace_mode(paragraphs_data)

        # 设置中文字体
        self._set_chinese_font(doc, target_lang)

        # 保存结果
        result_path = self._generate_result_path(source_path, target_lang)
        doc.save(result_path)

        return result_path

    def _write_replace_mode(self, paragraphs_data):
        """
        替换模式：直接替换原文

        Args:
            paragraphs_data: 段落数据列表
        """
        for item in paragraphs_data:
            paragraph = item['element']
            translated_text = item['translated']

            if translated_text and paragraph.runs:
                # 清空段落内容
                for run in paragraph.runs:
                    run.text = ""
                # 在第一个 run 中写入翻译文本
                paragraph.runs[0].text = translated_text

    def _write_parallel_mode(self, paragraphs_data, target_lang):
        """
        对照模式：保留原文，添加译文在下
        译文使用虚线下划线样式

        Args:
            paragraphs_data: 段落数据列表
            target_lang: 目标语言
        """
        import copy
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 需要按照段落/表格单元格分组处理
        processed_paragraphs = set()

        for item in paragraphs_data:
            paragraph = item['element']

            # 避免重复处理同一个段落
            if id(paragraph) in processed_paragraphs:
                continue
            processed_paragraphs.add(id(paragraph))

            original_text = item['text']
            translated_text = item['translated']

            if not translated_text:
                continue

            # 获取父元素（文档主体或表格单元格）
            parent = paragraph._element.getparent()

            # 创建新的段落元素（使用 deepcopy 复制）
            new_para_element = copy.deepcopy(paragraph._element)

            # 在原文段落后插入新段落
            parent.insert(
                parent.index(paragraph._element) + 1,
                new_para_element
            )

            # 创建新的段落对象
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_para_element, paragraph._parent)

            # 设置译文内容
            if new_paragraph.runs:
                # 清空新段落
                for run in new_paragraph.runs:
                    run.text = ""
                # 在第一个 run 中写入翻译文本
                new_paragraph.runs[0].text = translated_text

                # 设置译文样式：虚线下划线
                self._set_translation_style(new_paragraph.runs[0], target_lang)

    def _set_translation_style(self, run, target_lang):
        """
        设置译文样式：虚线下划线、蓝色

        Args:
            run: docx run 对象
            target_lang: 目标语言
        """
        from docx.enum.text import WD_UNDERLINE

        # 设置下划线样式（虚线）
        run.font.underline = WD_UNDERLINE.DASH

        # 设置字体颜色（蓝色）
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)  # 道奇蓝

        # 设置字体大小略小
        from docx.shared import Pt
        run.font.size = Pt(10.5)  # 五号字稍小

        # 设置中文字体
        try:
            if run._element.rPr is not None:
                from docx.oxml.ns import qn
                rfonts = run._element.rPr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = run._element.rPr.makeelement(qn('w:rFonts'))
                    run._element.rPr.append(rfonts)

                if target_lang == 'zh':
                    rfonts.set(qn('w:eastAsia'), '微软雅黑')
                    rfonts.set(qn('w:ascii'), '微软雅黑')
                    rfonts.set(qn('w:hAnsi'), '微软雅黑')
                else:
                    rfonts.set(qn('w:eastAsia'), 'Arial')
        except Exception:
            pass

    def _read_runs(self, runs, texts):
        """
        读取 runs 中的文本

        Args:
            runs: python-docx 的 runs 对象
            texts: 存储文本的列表
        """
        for run in runs:
            text = run.text
            if text and text.strip():
                texts.append({
                    'text': text,
                    'translated': None,
                    'complete': False
                })

    def _write_runs(self, runs, texts, index):
        """
        将翻译结果写回 runs

        Args:
            runs: python-docx 的 runs 对象
            texts: 存储翻译结果的列表
            index: 当前处理到的文本索引

        Returns:
            int: 更新后的索引
        """
        for run in runs:
            text = run.text
            if text and text.strip():
                if index < len(texts):
                    item = texts[index]
                    print(f"DEBUG _write_runs: index={index}, original={repr(text[:50])}, translated_type={type(item['translated'])}, translated={repr(str(item['translated'])[:50])}")
                    if item['translated']:
                        run.text = item['translated']
                    index += 1
        return index

    def _set_chinese_font(self, doc, target_lang: str):
        """
        设置文档的中文字体，确保中文能正确显示

        Args:
            doc: python-docx 文档对象
            target_lang: 目标语言代码
        """
        # 为所有段落设置中文字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    # 检查 run 是否有 rPr 属性
                    if run._element.rPr is not None:
                        # 获取或创建 rFonts 元素
                        rfonts = run._element.rPr.find(qn('w:rFonts'))
                        if rfonts is None:
                            rfonts = run._element.rPr.makeelement(qn('w:rFonts'))
                            run._element.rPr.append(rfonts)

                        # 设置东亚字体（中日韩字符）
                        if target_lang == 'zh':
                            rfonts.set(qn('w:eastAsia'), '微软雅黑')
                            rfonts.set(qn('w:ascii'), '微软雅黑')
                            rfonts.set(qn('w:hAnsi'), '微软雅黑')
                        else:
                            rfonts.set(qn('w:eastAsia'), 'Arial')
                except Exception as e:
                    # 如果设置失败，跳过该 run
                    pass

        # 为表格中的文本设置字体
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                if run._element.rPr is not None:
                                    rfonts = run._element.rPr.find(qn('w:rFonts'))
                                    if rfonts is None:
                                        rfonts = run._element.rPr.makeelement(qn('w:rFonts'))
                                        run._element.rPr.append(rfonts)

                                    if target_lang == 'zh':
                                        rfonts.set(qn('w:eastAsia'), '微软雅黑')
                                        rfonts.set(qn('w:ascii'), '微软雅黑')
                                        rfonts.set(qn('w:hAnsi'), '微软雅黑')
                                    else:
                                        rfonts.set(qn('w:eastAsia'), 'Arial')
                            except Exception:
                                pass

    def _generate_result_path(self, source_path: str, target_lang: str = "en") -> str:
        """
        生成结果文件路径

        Args:
            source_path: 源文件路径
            target_lang: 目标语言代码（如 en, zh 等）

        Returns:
            str: 结果文件路径
        """
        source = Path(source_path)
        # 使用源文件的文件名 + 语言代码
        filename = f"{source.stem}_{target_lang}{source.suffix}"
        return str(settings.TRANSLATE_DIR / filename)
